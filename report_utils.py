import csv
import unicodedata
from docx import Document
from fpdf import FPDF
import pandas as pd


def strip_accents(text: str) -> str:
    """Usuń znaki diakrytyczne (ogonek) z tekstu."""
    return ''.join(
        c for c in unicodedata.normalize('NFD', text)
        if unicodedata.category(c) != 'Mn'
    )


def save_report_csv(filepath, analysis_results):
    with open(filepath, 'w', newline='', encoding='utf-8-sig') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow([
            'Plik', 'Liczba słów', 'Liczba unikalnych słów', 'Średnia długość słowa',
            'Liczba zdań', 'Średnia długość zdania', 'Liczba znaków',
            'Najdłuższe słowo', 'Długość najdłuższego', 'Najkrótsze słowo', 'Długość najkrótszego'
        ])

        for r in analysis_results:
            writer.writerow([
                r['filename'], r['total_words'], r['unique_words'], f"{r['avg_word_len']:.2f}",
                r['sentence_count'], f"{r['avg_sentence_len']:.2f}", r['char_count'],
                r['longest_word'], r['longest_len'], r['shortest_word'], r['shortest_len']
            ])

            writer.writerow([])
            writer.writerow(['Top 10 najczęstszych słów'])
            writer.writerow(['Słowo', 'Liczba'])
            for word, count in r['top_words']:
                writer.writerow([word, count])
            writer.writerow([])


def save_report_docx(filepath, analysis_results):
    doc = Document()

    for r in analysis_results:
        doc.add_heading(f"Analiza pliku: {r['filename']}", level=1)

        table = doc.add_table(rows=10, cols=2)
        table.style = 'Light List Accent 1'

        data = [
            ("Liczba wszystkich słów", r['total_words']),
            ("Liczba unikalnych słów", r['unique_words']),
            ("Średnia długość słowa", f"{r['avg_word_len']:.2f}"),
            ("Liczba zdań", r['sentence_count']),
            ("Średnia długość zdania", f"{r['avg_sentence_len']:.2f}"),
            ("Liczba znaków", r['char_count']),
            ("Najdłuższe słowo", r['longest_word']),
            ("Długość najdłuższego", r['longest_len']),
            ("Najkrótsze słowo", r['shortest_word']),
            ("Długość najkrótszego", r['shortest_len']),
        ]

        for i, (label, val) in enumerate(data):
            row = table.rows[i].cells
            row[0].text = label
            row[1].text = str(val)

        doc.add_paragraph("Top 10 najczęstszych słów:")
        top_table = doc.add_table(rows=1, cols=2)
        top_table.style = 'Light List Accent 2'
        hdr = top_table.rows[0].cells
        hdr[0].text = 'Słowo'
        hdr[1].text = 'Liczba'

        for word, count in r['top_words']:
            row_cells = top_table.add_row().cells
            row_cells[0].text = word
            row_cells[1].text = str(count)

        doc.add_paragraph()  # odstęp

    doc.save(filepath)


def save_report_pdf(filepath, analysis_results):
    pdf = FPDF()
    pdf.add_page()

    # Dodanie czcionek z unicode
    pdf.add_font('DejaVu', '', 'fonts/DejaVuSans.ttf', uni=True)
    pdf.add_font('DejaVu', 'B', 'fonts/DejaVuSans-Bold.ttf', uni=True)

    # Tytuł
    pdf.set_font('DejaVu', 'B', 16)
    pdf.cell(0, 10, 'Raport analizy tekstu', ln=True, align='C')
    pdf.ln(10)

    # Ustawienia tabel
    col_widths = [80, 30]
    row_height = 8

    for r in analysis_results:
        # Nagłówek rozdziału
        pdf.set_font('DejaVu', 'B', 14)
        pdf.cell(0, 10, f"Analiza pliku: {r['filename']}", ln=True)
        pdf.ln(2)

        # Tabela statystyk
        pdf.set_font('DejaVu', '', 12)
        labels = [
            "Liczba wszystkich słów", "Liczba unikalnych słów", "Średnia długość słowa",
            "Liczba zdań", "Średnia długość zdania", "Liczba znaków",
            "Najdłuższe słowo", "Długość najdłuższego",
            "Najkrótsze słowo", "Długość najkrótszego"
        ]
        values = [
            r['total_words'], r['unique_words'], f"{r['avg_word_len']:.2f}",
            r['sentence_count'], f"{r['avg_sentence_len']:.2f}", r['char_count'],
            r['longest_word'], r['longest_len'], r['shortest_word'], r['shortest_len']
        ]
        # Rysowanie tabeli
        for label, val in zip(labels, values):
            pdf.cell(col_widths[0], row_height, label, border=1)
            pdf.cell(col_widths[1], row_height, str(val), border=1, ln=True)
        pdf.ln(5)

        # Tabela Top10
        pdf.set_font('DejaVu', 'B', 12)
        pdf.cell(col_widths[0], row_height, 'Słowo', border=1)
        pdf.cell(col_widths[1], row_height, 'Liczba', border=1, ln=True)
        pdf.set_font('DejaVu', '', 12)
        for word, count in r['top_words']:
            pdf.cell(col_widths[0], row_height, word, border=1)
            pdf.cell(col_widths[1], row_height, str(count), border=1, ln=True)
        pdf.ln(10)

    pdf.output(filepath)


def save_report_xlsx(filepath, analysis_results):
    """
    Tworzy plik .xlsx z:
    - arkuszem 'Stats' z podsumowaniem (jeden wiersz = jeden plik)
    - arkuszem 'Top10' z oddzielnymi tabelami top10 słów dla każdego pliku obok siebie
    """
    stats_rows = []
    for r in analysis_results:
        stats_rows.append({
            'Plik': r['filename'],
            'Liczba słów': r['total_words'],
            'Liczba unikalnych słów': r['unique_words'],
            'Średnia długość słowa': round(r['avg_word_len'], 2),
            'Liczba zdań': r['sentence_count'],
            'Średnia długość zdania': round(r['avg_sentence_len'], 2),
            'Liczba znaków': r['char_count'],
            'Najdłuższe słowo': r['longest_word'],
            'Długość najdłuższego': r['longest_len'],
            'Najkrótsze słowo': r['shortest_word'],
            'Długość najkrótszego': r['shortest_len'],
        })
    df_stats = pd.DataFrame(stats_rows)

    with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
        df_stats.to_excel(writer, sheet_name='Stats', index=False)
        workbook = writer.book
        ws = workbook.create_sheet('Top10')

        start_col = 1
        for r in analysis_results:
            ws.cell(row=1, column=start_col, value=f"Top 10: {r['filename']}")
            ws.cell(row=2, column=start_col, value='Słowo')
            ws.cell(row=2, column=start_col+1, value='Liczba')
            for idx, (word, count) in enumerate(r['top_words'], start=1):
                ws.cell(row=2+idx, column=start_col, value=word)
                ws.cell(row=2+idx, column=start_col+1, value=count)
            start_col += 3
